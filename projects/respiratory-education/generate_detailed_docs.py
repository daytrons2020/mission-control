#!/usr/bin/env python3
"""
Create detailed Word documents from outline markdown files.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def create_detailed_word_doc(input_path, output_path, title):
    """Create a comprehensive Word document from outline content."""
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Add title page
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("\nComprehensive Educational Guide\nRespiratory Care Department")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x5A, 0xA8)
    run.font.italic = True
    
    doc.add_page_break()
    
    # Process content
    lines = content.split('\n')
    in_speaker_notes = False
    current_section = None
    
    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        
        # Skip marp directives
        if line.startswith('---') or line.startswith('marp:') or line.startswith('theme:') or line.startswith('paginate:') or line.startswith('backgroundColor:'):
            continue
        if line.startswith('<!--') or line.endswith('-->'):
            continue
        
        # Main title (H1)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            in_speaker_notes = False
        
        # Section headers (H2)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            in_speaker_notes = False
        
        # Subsection headers (H3)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            in_speaker_notes = False
        
        # Sub-subsection (H4)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            in_speaker_notes = False
        
        # Speaker notes section
        elif '**Speaker Notes:**' in line:
            doc.add_heading('Speaker Notes', level=4)
            in_speaker_notes = True
        
        # Table handling
        elif '|' in line and not line.startswith('|--'):
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                p = doc.add_paragraph()
                p.style = 'Intense Quote'
                run = p.add_run(' | '.join(cells))
                run.font.size = Pt(10)
        
        # Regular content
        else:
            # Clean markdown formatting
            clean_line = line
            clean_line = re.sub(r'\*\*\*([^*]+)\*\*\*', r'\1', clean_line)  # Bold+italic
            clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_line)      # Bold
            clean_line = re.sub(r'\*([^*]+)\*', r'\1', clean_line)          # Italic
            clean_line = re.sub(r'`([^`]+)`', r'\1', clean_line)            # Code
            clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)  # Links
            
            if clean_line.strip():
                # Bullet points
                if clean_line.strip().startswith('- ') or clean_line.strip().startswith('* '):
                    text = clean_line.strip()[2:]
                    p = doc.add_paragraph(text, style='List Bullet')
                    if in_speaker_notes:
                        p.paragraph_format.left_indent = Inches(0.5)
                # Numbered lists
                elif re.match(r'^\d+\.\s', clean_line.strip()):
                    text = re.sub(r'^\d+\.\s', '', clean_line.strip())
                    p = doc.add_paragraph(text, style='List Number')
                # Regular paragraphs
                else:
                    p = doc.add_paragraph(clean_line)
                    if in_speaker_notes:
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.space_after = Pt(6)
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Document —")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    
    doc.save(output_path)
    print(f"Created detailed Word document: {output_path}")


def main():
    templates_dir = "/root/.openclaw/workspace/projects/respiratory-education/templates"
    output_dir = "/root/.openclaw/workspace/projects/respiratory-education/output"
    
    os.makedirs(output_dir, exist_ok=True)
    
    outline_files = [
        ("01-anatomy-physiology-outline.md", "01-Anatomy-Physiology-Detailed.docx", "Respiratory Anatomy and Physiology"),
        ("02-copd-management-outline.md", "02-COPD-Management-Detailed.docx", "COPD Management"),
        ("03-mechanical-ventilation-outline.md", "03-Mechanical-Ventilation-Detailed.docx", "Mechanical Ventilation Management"),
        ("04-patient-assessment-outline.md", "04-Patient-Assessment-Detailed.docx", "Patient Assessment in Respiratory Care"),
        ("05-emergency-procedures-outline.md", "05-Emergency-Procedures-Detailed.docx", "Emergency Respiratory Procedures"),
    ]
    
    for md_file, docx_file, title in outline_files:
        input_path = os.path.join(templates_dir, md_file)
        output_path = os.path.join(output_dir, docx_file)
        
        if os.path.exists(input_path):
            print(f"\nProcessing {md_file}...")
            create_detailed_word_doc(input_path, output_path, title)
        else:
            print(f"Warning: {input_path} not found")
    
    print("\n✓ All detailed Word documents created!")


if __name__ == "__main__":
    main()
